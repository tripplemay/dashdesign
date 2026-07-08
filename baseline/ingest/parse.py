"""Parse uploaded documents (PDF / DOCX / TXT) into sectioned text.

Uses permissive-licensed libraries only (pypdfium2 = Apache-2.0/BSD,
python-docx = MIT) so the closed-source desktop app can bundle them. Section
anchors (headings / page numbers) are preserved so extracted quotes can be
traced back to ``evidence_index.section``. Parsing runs locally — sensitive
B-side originals never leave the machine; only extracted text is sent to the
LLM gateway.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}


@dataclass(frozen=True)
class ParsedSection:
    section: str  # heading or "第 N 页"
    text: str


@dataclass(frozen=True)
class ParsedDocument:
    file_name: str
    suffix: str
    sections: List[ParsedSection] = field(default_factory=list)

    def full_text(self) -> str:
        return "\n\n".join(f"【{s.section}】\n{s.text}" for s in self.sections)


class UnsupportedDocument(Exception):
    pass


def parse_document(path: Path) -> ParsedDocument:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        sections = _parse_pdf(path)
    elif suffix == ".docx":
        sections = _parse_docx(path)
    elif suffix in (".txt", ".md"):
        sections = _parse_text(path)
    elif suffix == ".doc":
        raise UnsupportedDocument(
            "旧版 .doc 二进制格式无法直接解析，请提供同名 .pdf/.docx，或先转换为 .docx。"
        )
    else:
        raise UnsupportedDocument(f"不支持的文档类型：{suffix or path.name}")
    return ParsedDocument(file_name=path.name, suffix=suffix, sections=sections)


def _parse_pdf(path: Path) -> List[ParsedSection]:
    import pypdfium2 as pdfium

    sections: List[ParsedSection] = []
    pdf = pdfium.PdfDocument(str(path))
    try:
        for index in range(len(pdf)):
            page = pdf[index]
            textpage = page.get_textpage()
            text = textpage.get_text_range().strip()
            textpage.close()
            page.close()
            if text:
                sections.append(ParsedSection(section=f"第 {index + 1} 页", text=text))
    finally:
        pdf.close()
    return sections


def _parse_docx(path: Path) -> List[ParsedSection]:
    import docx

    document = docx.Document(str(path))
    sections: List[ParsedSection] = []
    current_heading = "正文"
    buffer: List[str] = []

    def flush() -> None:
        if buffer:
            sections.append(ParsedSection(section=current_heading, text="\n".join(buffer).strip()))
            buffer.clear()

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading") or style.startswith("title"):
            flush()
            current_heading = text
        else:
            buffer.append(text)
    flush()
    # 表格内容并入正文尾部
    table_lines: List[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
    if table_lines:
        sections.append(ParsedSection(section="表格", text="\n".join(table_lines)))
    return sections


def _parse_text(path: Path) -> List[ParsedSection]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    return [ParsedSection(section=path.name, text=text)] if text else []
