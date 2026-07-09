"""Unit tests for document ingestion, grounded extraction, and merge."""

from __future__ import annotations

import copy
import json
from pathlib import Path

import pytest

from app_runtime import baseline_path
from baseline import merge
from baseline.extract import extract_candidates
from baseline.ingest.parse import UnsupportedDocument, parse_document
from baseline.schema import validation_errors


@pytest.fixture(scope="module")
def bundled_baseline() -> dict:
    return json.loads(baseline_path().read_text(encoding="utf-8"))


# --- parsing ---------------------------------------------------------
class TestParse:
    def test_txt(self, tmp_path: Path) -> None:
        p = tmp_path / "intro.txt"
        p.write_text("项目介绍正文", encoding="utf-8")
        doc = parse_document(p)
        assert doc.suffix == ".txt"
        assert "项目介绍正文" in doc.full_text()

    def test_docx_headings_become_sections(self, tmp_path: Path) -> None:
        import docx

        d = docx.Document()
        d.add_heading("一、项目概述", level=1)
        d.add_paragraph("面向青少年的 AI 数字化创作课程。")
        d.add_heading("二、课程体系", level=1)
        d.add_paragraph("AI 绘图、AI 视频、AI 漫剧。")
        path = tmp_path / "intro.docx"
        d.save(str(path))
        doc = parse_document(path)
        headings = [s.section for s in doc.sections]
        assert "一、项目概述" in headings and "二、课程体系" in headings
        assert "AI 绘图" in doc.full_text()

    def test_legacy_doc_rejected(self, tmp_path: Path) -> None:
        p = tmp_path / "old.doc"
        p.write_bytes(b"\xd0\xcf\x11\xe0")  # OLE2 magic
        with pytest.raises(UnsupportedDocument, match="doc"):
            parse_document(p)


# --- extraction grounding -------------------------------------------
class TestExtract:
    def _parsed(self, tmp_path: Path, text: str):
        p = tmp_path / "src.txt"
        p.write_text(text, encoding="utf-8")
        return parse_document(p)

    def test_ungrounded_candidates_dropped(self, tmp_path: Path, bundled_baseline: dict) -> None:
        parsed = self._parsed(tmp_path, "孩子在 AI 数字化创作中提升想象力与表达力。")

        def fake_chat(messages, response_format):  # noqa: ANN001
            return json.dumps(
                {
                    "source_context_hint": "to_c_marketing_docs",
                    "evidence": [
                        {"id": "ev_new_1", "section": "src.txt", "quote": "提升想象力与表达力"},
                        {"id": "ev_bogus", "section": "src.txt", "quote": "保证考上名校"},
                    ],
                    "candidates": [
                        {
                            "target": "consumer_baseline.parent_value",
                            "text": "提升孩子的想象力与表达力",
                            "confidence": 0.8,
                            "evidence": ["ev_new_1"],
                        },
                        {
                            "target": "consumer_baseline.core_messages",
                            "text": "保证考上名校",
                            "confidence": 0.9,
                            "evidence": ["ev_bogus"],  # 引文不在原文 → 应被丢弃
                        },
                    ],
                },
                ensure_ascii=False,
            )

        result = extract_candidates(parsed, bundled_baseline, fake_chat)
        texts = [c["text"] for c in result["candidates"]]
        assert "提升孩子的想象力与表达力" in texts
        assert "保证考上名校" not in texts  # 幻觉引文被接地校验拦掉
        assert result["document"]["document_id"]

    def test_nonstandard_shape_is_salvaged(self, tmp_path: Path, bundled_baseline: dict) -> None:
        # 复现网关模型不遵守 json_schema、自创字段名的情形，应被 _adapt 抢救
        parsed = self._parsed(tmp_path, "商业皮肤学课程产品化方案，面向美业代理的成交系统。")

        def fake_chat(messages, response_format=None):  # noqa: ANN001
            return json.dumps({
                "project": "商业皮肤学",
                "audience_mode": "to_b_partnership",
                "consumer_baseline": {"titles": [], "subtitles": []},
                "source_facts": {
                    "consumer_safe_facts": [],
                    "business_terms": [
                        {"type": "topic", "text": "商业皮肤学课程产品化方案",
                         "evidence": {"quote": "商业皮肤学课程产品化方案"}}
                    ],
                },
            }, ensure_ascii=False)

        result = extract_candidates(parsed, bundled_baseline, fake_chat)
        assert len(result["candidates"]) == 1
        assert result["candidates"][0]["target"] == "source_facts.business_terms"
        assert result["candidates"][0]["text"] == "商业皮肤学课程产品化方案"

    def test_whitespace_quote_not_grounded(self, tmp_path: Path, bundled_baseline: dict) -> None:
        parsed = self._parsed(tmp_path, "孩子提升表达力。")

        def fake_chat(messages, response_format):  # noqa: ANN001
            return json.dumps({
                "source_context_hint": "manual",
                "evidence": [{"id": "ev_ws", "section": "src.txt", "quote": "   "}],
                "candidates": [{"target": "consumer_baseline.core_messages",
                                "text": "编造内容", "confidence": 0.9, "evidence": ["ev_ws"]}],
            }, ensure_ascii=False)

        result = extract_candidates(parsed, bundled_baseline, fake_chat)
        assert result["candidates"] == []  # 纯空白 quote 不算接地


# --- merge classification + apply -----------------------------------
class TestMerge:
    def _extraction(self, blocked_kw: str) -> dict:
        return {
            "source_context_hint": "mixed_docs",
            "document": {
                "document_id": "new_doc_2026",
                "file": "new.pdf",
                "type": "pdf",
                "role": "supplementary_document",
                "status": "active",
            },
            "evidence": [
                {"id": "ev_x_ok", "document_id": "new_doc_2026", "section": "一", "quote": "想象力"},
                {"id": "ev_x_b", "document_id": "new_doc_2026", "section": "二", "quote": "加盟"},
            ],
            "candidates": [
                {"target": "consumer_baseline.parent_value", "text": "培养孩子的想象力",
                 "confidence": 0.82, "evidence": ["ev_x_ok"]},
                {"target": "consumer_baseline.core_messages", "text": f"欢迎{blocked_kw}",
                 "confidence": 0.9, "evidence": ["ev_x_ok"]},
                {"target": "consumer_baseline.student_value", "text": "保证掌握全部技能",
                 "confidence": 0.95, "evidence": ["ev_x_ok"]},
                {"target": "consumer_baseline.parent_value", "text": "也许有点用",
                 "confidence": 0.3, "evidence": ["ev_x_ok"]},
                {"target": "source_facts.business_terms", "text": f"{blocked_kw}扶持政策",
                 "confidence": 0.9, "evidence": ["ev_x_b"]},
            ],
        }

    def test_classification(self, bundled_baseline: dict) -> None:
        blocked_kw = bundled_baseline["consumer_baseline"]["blocked_keywords"][0]
        report = merge.build_merge_report(bundled_baseline, self._extraction(blocked_kw))
        by_text = {c.text: c for c in report.changes}
        assert by_text["培养孩子的想象力"].governance == merge.OK
        assert by_text["培养孩子的想象力"].accepted is True
        assert by_text[f"欢迎{blocked_kw}"].governance == merge.BLOCKED
        assert by_text[f"欢迎{blocked_kw}"].accepted is False
        assert by_text["保证掌握全部技能"].governance == merge.FORBIDDEN
        assert by_text["保证掌握全部技能"].accepted is False
        assert by_text["也许有点用"].governance == merge.LOW_CONF
        assert by_text["也许有点用"].accepted is False
        assert by_text[f"{blocked_kw}扶持政策"].governance == merge.B_SIDE
        assert by_text[f"{blocked_kw}扶持政策"].accepted is True  # B 端桶允许

    def test_duplicate_skipped(self, bundled_baseline: dict) -> None:
        existing = bundled_baseline["consumer_baseline"]["parent_value"][0]["text"]
        extraction = {
            "document": {}, "evidence": [], "source_context_hint": "manual",
            "candidates": [{"target": "consumer_baseline.parent_value", "text": existing,
                            "confidence": 0.9, "evidence": []}],
        }
        report = merge.build_merge_report(bundled_baseline, extraction)
        assert report.changes == []

    def test_apply_produces_valid_draft(self, bundled_baseline: dict) -> None:
        blocked_kw = bundled_baseline["consumer_baseline"]["blocked_keywords"][0]
        report = merge.build_merge_report(bundled_baseline, self._extraction(blocked_kw))
        draft = merge.apply_report(bundled_baseline, report, "2026.07.09", [bundled_baseline["version"]])
        # 结构合法
        assert validation_errors(draft) == []
        # 追加式 + 版本链
        assert draft["status"] == "draft"
        assert draft["parent_version"] == bundled_baseline["version"]
        assert draft["version"] == "2026.07.09.1"
        # 只有被采纳的进了 C 端；被拦截/需审的没进
        pv_texts = [i["text"] for i in draft["consumer_baseline"]["parent_value"]]
        assert "培养孩子的想象力" in pv_texts
        cm_texts = [i["text"] for i in draft["consumer_baseline"]["core_messages"]]
        assert f"欢迎{blocked_kw}" not in cm_texts
        sv_texts = [i["text"] for i in draft["consumer_baseline"]["student_value"]]
        assert "保证掌握全部技能" not in sv_texts
        # B 端事实进了 business_terms
        bt_texts = [i["text"] for i in draft["source_facts"]["business_terms"]]
        assert f"{blocked_kw}扶持政策" in bt_texts
        # 新文档 + 源上下文
        assert any(d["document_id"] == "new_doc_2026" for d in draft["source_documents"])
        assert draft["source_context"] == "mixed_docs"

    def test_apply_only_accepted(self, bundled_baseline: dict) -> None:
        blocked_kw = bundled_baseline["consumer_baseline"]["blocked_keywords"][0]
        report = merge.build_merge_report(bundled_baseline, self._extraction(blocked_kw))
        # 人工把"需人工"的承诺项也批准（模拟审校采纳）→ 应能进入草稿
        for c in report.changes:
            if c.text == "保证掌握全部技能":
                c.accepted = True
        draft = merge.apply_report(bundled_baseline, report, "2026.07.09", [bundled_baseline["version"]])
        sv_texts = [i["text"] for i in draft["consumer_baseline"]["student_value"]]
        assert "保证掌握全部技能" in sv_texts  # 人工采纳后进入
